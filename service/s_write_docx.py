import os
from service.s_folder_name import FolderNameService
from service.s_folders import FoldersService
from docx import Document
from docx.shared import Inches


class WriteDocxService:
    EXTENSIONS = [".jpg", ".png", ".jpeg", ".gif", ".bmp", ".tiff", ".ico", ".webp"]

    def __init__(self, output_path='qa_evidence.docx'):
        self.document = Document()
        self.output_path = output_path

    def write_docx(self, pasta_selecionada):
        self.write_image_to_docx(pasta_selecionada)
        self.document.save(self.output_path)
        return "Documento gerado com sucesso"

    def write_image_to_docx(self, target_folder):
        folders = FoldersService.get_folders(target_folder)
        files = FoldersService.get_files_with_extension(target_folder, self.EXTENSIONS)

        if files:
            folder_name = FolderNameService.get_folder_name(target_folder)
            self.document.add_heading(folder_name, level=1)
            for file in files:
                file_path = os.path.join(target_folder, file)
                self.document.add_paragraph(file)
                picture = self.document.add_picture(file_path)
                picture.width = Inches(4.8)
                picture.height = Inches(3.0)
                self.document.add_paragraph()  # Adiciona uma linha em branco após a imagem

        for folder in folders:
            subfolder_path = os.path.join(target_folder, folder)
            self.write_image_to_docx(subfolder_path)